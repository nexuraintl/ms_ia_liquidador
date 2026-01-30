"""
EXTRACTOR DE DOCUMENTOS CON GUARDADO AUTOMÁTICO - VERSIÓN CORREGIDA
==================================================================

Módulo para extraer texto de diferentes tipos de archivos:
- PDFs (con conversión a imagen para OCR fallback)
- Imágenes (JPG, PNG) usando Google Vision OCR
- Excel (XLSX, XLS)
- Word (DOCX, DOC)

CORRECCIÓN: PDF → Imagen → OCR para casos de poco texto extraído.

Autor: Miguel Angel Jaramillo Durango
"""

import os
import io
import json
import logging
import asyncio
from pathlib import Path
from typing import Dict, Any
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor
from functools import partial
import xml.etree.ElementTree as ET

# Procesamiento de archivos
import PyPDF2
from PIL import Image
import pandas as pd
from docx import Document

# Procesamiento de emails
import email
from email.utils import parsedate_to_datetime, parseaddr
from email.header import decode_header
try:
    import extract_msg
    EXTRACT_MSG_DISPONIBLE = True
except ImportError:
    EXTRACT_MSG_DISPONIBLE = False
    logging.warning("extract-msg no disponible - archivos .msg limitados")

# Nuevas dependencias para PDF → Imagen
try:
    import pdf2image
    PDF2IMAGE_DISPONIBLE = True
except ImportError:
    PDF2IMAGE_DISPONIBLE = False
    logging.warning("pdf2image no disponible - OCR fallback para PDF limitado")

try:
    import fitz  # PyMuPDF
    PYMUPDF_DISPONIBLE = True
except ImportError:
    PYMUPDF_DISPONIBLE = False

# PDF Plumber para extracción mejorada
try:
    import pdfplumber
    PDFPLUMBER_DISPONIBLE = True
except ImportError:
    PDFPLUMBER_DISPONIBLE = False
    logging.warning("pdfplumber no disponible - extracción PDF limitada")

# Google Vision para OCR
from google.cloud import vision

# FastAPI
from fastapi import UploadFile

# Configuración de logging
logger = logging.getLogger(__name__)

# ===============================
# PROCESADOR DE ARCHIVOS CON GUARDADO AUTOMÁTICO
# ===============================

class ProcesadorArchivos:
    """
    Extrae texto de diferentes tipos de archivos usando las mejores técnicas
    disponibles para cada formato.
    
    """
    
    def __init__(self):
        """Inicializa el procesador con configuración de OCR y carpetas de guardado"""
        self.vision_client = self._configurar_vision()
        self._crear_carpetas_guardado()
        self._verificar_dependencias_pdf()
        logger.info("ProcesadorArchivos inicializado con guardado automático")
    
    def _verificar_dependencias_pdf(self):
        """Verifica y reporta las dependencias disponibles para extracción PDF y conversión a imagen"""
        # Verificar PDF Plumber (método principal)
        if PDFPLUMBER_DISPONIBLE:
            logger.info(" pdfplumber disponible para extracción principal de PDF")
        else:
            logger.warning(" pdfplumber no disponible. Usando PyPDF2 como principal")
            logger.warning("   Instala: pip install pdfplumber")
        
        # Verificar PyPDF2 (fallback)
        logger.info(" PyPDF2 disponible como fallback")
        
        # Verificar conversión PDF → Imagen para OCR
        if PDF2IMAGE_DISPONIBLE:
            logger.info("pdf2image disponible para conversión PDF → Imagen")
        elif PYMUPDF_DISPONIBLE:
            logger.info(" PyMuPDF disponible para conversión PDF → Imagen")
        else:
            logger.warning(" Sin dependencias para PDF → Imagen. OCR fallback limitado")
            logger.warning("   Instala: pip install pdf2image PyMuPDF")
        
        # Verificar dependencias de email
        if EXTRACT_MSG_DISPONIBLE:
            logger.info(" extract-msg disponible para archivos .msg")
        else:
            logger.warning(" extract-msg no disponible. Archivos .msg limitados")
            logger.warning("   Instala: pip install extract-msg")
    
    def _configurar_vision(self):
        """
        Configura Google Vision para OCR (opcional).
        
        Returns:
            vision.ImageAnnotatorClient o None si no está configurado
        """
        try:
            from dotenv import load_dotenv
            load_dotenv()
            
            credentials_path = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
            if credentials_path and os.path.exists(credentials_path):
                vision_client = vision.ImageAnnotatorClient()
                logger.info("Google Vision configurado correctamente para OCR")
                return vision_client
            else:
                logger.warning("Google Vision no configurado - OCR limitado")
                return None
                
        except Exception as e:
            logger.warning(f"No se pudo configurar Google Vision: {e}")
            return None
    
    def _crear_carpetas_guardado(self):
        """Crea las carpetas necesarias para guardar extracciones"""
        try:
            # Crear carpeta base de extracciones
            self.carpeta_base = Path("Results/Extracciones")
            self.carpeta_base.mkdir(parents=True, exist_ok=True)
            
            # Crear carpeta por fecha actual
            fecha_hoy = datetime.now().strftime("%Y-%m-%d")
            self.carpeta_fecha = self.carpeta_base / fecha_hoy
            self.carpeta_fecha.mkdir(parents=True, exist_ok=True)
            
            logger.info(f"Carpetas de guardado creadas: {self.carpeta_fecha}")
            
        except Exception as e:
            logger.error(f"Error creando carpetas de guardado: {e}")
            # Fallback a carpeta actual
            self.carpeta_fecha = Path(".")
    
    def _guardar_texto_extraido(self, nombre_archivo: str, texto_extraido: str, 
                               metodo_extraccion: str, metadatos: Dict = None) -> str:
        """
        Guarda el texto extraído en un archivo organizado.
        
        Args:
            nombre_archivo: Nombre del archivo original
            texto_extraido: Texto que se extrajo
            metodo_extraccion: Método usado (PDF, OCR, EXCEL, WORD)
            metadatos: Información adicional sobre la extracción
            
        Returns:
            str: Ruta del archivo guardado
        """
        try:
            # Crear timestamp único
            timestamp = datetime.now().strftime("%H%M%S")
            
            # Limpiar nombre de archivo para usar como nombre base
            nombre_base = "".join(c for c in nombre_archivo if c.isalnum() or c in "._-")
            if len(nombre_base) > 50:  # Limitar longitud
                nombre_base = nombre_base[:50]
            
            # Crear nombre de archivo único
            nombre_salida = f"{timestamp}_{metodo_extraccion}_{nombre_base}.txt"
            ruta_archivo = self.carpeta_fecha / nombre_salida
            
            # Preparar contenido completo
            contenido_completo = f"""EXTRACCIÓN DE TEXTO - PRELIQUIDADOR v2.0
============================================

INFORMACIÓN DEL ARCHIVO:
- Archivo original: {nombre_archivo}
- Método de extracción: {metodo_extraccion}
- Fecha y hora: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
- Caracteres extraídos: {len(texto_extraido)}

METADATOS ADICIONALES:
{json.dumps(metadatos or {}, indent=2, ensure_ascii=False)}

============================================
TEXTO EXTRAÍDO:
============================================

{texto_extraido}

============================================
FIN DE LA EXTRACCIÓN
============================================
"""
            
            # Guardar archivo
            with open(ruta_archivo, 'w', encoding='utf-8') as f:
                f.write(contenido_completo)
            
            logger.info(f" Texto extraído guardado: {ruta_archivo}")
            logger.info(f" Caracteres extraídos: {len(texto_extraido)}")
            
            return str(ruta_archivo)
            
        except Exception as e:
            logger.error(f" Error guardando texto extraído: {e}")
            return f"Error guardando: {str(e)}"
    
    def _validar_pdf(self, contenido_pdf: bytes, nombre_archivo: str) -> Dict[str, Any]:
        """
        Valida un PDF antes de intentar conversión a imagen.
        
        Args:
            contenido_pdf: Contenido binario del PDF
            nombre_archivo: Nombre del archivo para logging
            
        Returns:
            Dict con información de validación
        """
        validacion = {
            "valido": False,
            "error": None,
            "info": {},
            "metodo_recomendado": None
        }
        
        try:
            # Verificar tamaño mínimo
            if len(contenido_pdf) < 100:
                validacion["error"] = f"PDF demasiado pequeño: {len(contenido_pdf)} bytes"
                return validacion
            
            # Verificar header PDF
            if not contenido_pdf.startswith(b'%PDF-'):
                validacion["error"] = "No es un archivo PDF válido (header incorrecto)"
                return validacion
            
            # Intentar abrir con PyPDF2 para validación básica
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(contenido_pdf))
                num_paginas = len(pdf_reader.pages)
                
                validacion["info"]["paginas"] = num_paginas
                validacion["info"]["tamaño_bytes"] = len(contenido_pdf)
                
                if num_paginas == 0:
                    validacion["error"] = "PDF sin páginas"
                    return validacion
                
                # Verificar si está protegido
                if pdf_reader.is_encrypted:
                    validacion["error"] = "PDF protegido con contraseña"
                    return validacion
                
                validacion["valido"] = True
                
                # Recomendar método según dependencias disponibles
                if PDF2IMAGE_DISPONIBLE:
                    validacion["metodo_recomendado"] = "pdf2image"
                elif PYMUPDF_DISPONIBLE:
                    validacion["metodo_recomendado"] = "PyMuPDF"
                else:
                    validacion["error"] = "No hay librerías de conversión instaladas"
                    validacion["valido"] = False
                
            except Exception as e:
                validacion["error"] = f"PDF corrupto o inválido: {str(e)}"
                return validacion
                
        except Exception as e:
            validacion["error"] = f"Error validando PDF: {str(e)}"
        
        return validacion
    
    def _convertir_pdf_a_imagenes(self, contenido_pdf: bytes, nombre_archivo: str) -> list:
        """
        Convierte páginas de PDF a imágenes para OCR.
        
        Args:
            contenido_pdf: Contenido binario del PDF
            nombre_archivo: Nombre del archivo para logging
            
        Returns:
            list: Lista de imágenes en bytes (JPEG)
        """
        imagenes = []
        
        try:
            # Validar que el contenido no esté vacío
            if not contenido_pdf or len(contenido_pdf) < 100:
                logger.error(f" PDF demasiado pequeño o vacío: {len(contenido_pdf)} bytes")
                return []
            
            # Intentar con pdf2image primero
            if PDF2IMAGE_DISPONIBLE:
                try:
                    from pdf2image import convert_from_bytes
                    
                    logger.info(f" Convirtiendo PDF a imágenes con pdf2image: {nombre_archivo}")
                    
                    # Convertir PDF a imágenes con configuración robusta
                    pages = convert_from_bytes(
                        contenido_pdf,
                        dpi=300,  # Alta calidad para OCR
                        fmt='JPEG',
                        thread_count=1,  # Evitar problemas de concurrencia
                        first_page=1,
                        last_page=1000  # Limitar a 1000 páginas máximo
                    )
                    
                    # Convertir cada página a bytes
                    for i, page in enumerate(pages):
                        img_byte_arr = io.BytesIO()
                        page.save(img_byte_arr, format='JPEG', quality=95)
                        imagenes.append(img_byte_arr.getvalue())
                    
                    logger.info(f" pdf2image: {len(imagenes)} páginas convertidas")
                    return imagenes
                    
                except Exception as e:
                    logger.error(f" Error con pdf2image: {e}")
                    # Continuar con PyMuPDF
            
            # Intentar con PyMuPDF como alternativa
            if PYMUPDF_DISPONIBLE:
                try:
                    logger.info(f" Convirtiendo PDF a imágenes con PyMuPDF: {nombre_archivo}")
                    
                    # Abrir PDF desde bytes
                    pdf_document = fitz.open(stream=contenido_pdf, filetype="pdf")
                    
                    # Verificar que el PDF se abrió correctamente
                    if pdf_document.page_count == 0:
                        logger.error(f" PDF sin páginas válidas")
                        pdf_document.close()
                        return []
                    
                    # Limitar número de páginas
                    max_pages = min(pdf_document.page_count, 1000)
                    logger.info(f" Procesando {max_pages} de {pdf_document.page_count} páginas")
                    
                    # Convertir cada página
                    for page_num in range(max_pages):
                        page = pdf_document.load_page(page_num)
                        
                        # Renderizar página como imagen
                        mat = fitz.Matrix(3.0, 3.0)  # Factor de escala para calidad
                        pix = page.get_pixmap(matrix=mat)
                        
                        # Convertir a bytes JPEG
                        img_data = pix.tobytes("jpeg")
                        imagenes.append(img_data)
                    
                    pdf_document.close()
                    logger.info(f"PyMuPDF: {len(imagenes)} páginas convertidas")
                    return imagenes
                    
                except Exception as e:
                    logger.error(f" Error con PyMuPDF: {e}")
            
            # Si llegamos aquí, ambos métodos fallaron
            logger.error(f" Todos los métodos de conversión PDF a imagen  fallaron")
            logger.error(f"   pdf2image disponible: {PDF2IMAGE_DISPONIBLE}")
            logger.error(f"   PyMuPDF disponible: {PYMUPDF_DISPONIBLE}")
            
            if not PDF2IMAGE_DISPONIBLE and not PYMUPDF_DISPONIBLE:
                logger.error(f"   CAUSA: No hay librerías de conversión instaladas")
                logger.error(f"   SOLUCIÓN: pip install pdf2image PyMuPDF")
            
            return []
                
        except Exception as e:
            logger.error(f" Error general convirtiendo PDF a imágenes: {e}")
            logger.error(f"   Archivo: {nombre_archivo}")
            logger.error(f"   Tamaño: {len(contenido_pdf)} bytes")
            return []
    
    async def procesar_archivo(self, archivo: UploadFile) -> str:
        """
        Procesa un archivo y extrae su texto usando el método más apropiado.
        GUARDA AUTOMÁTICAMENTE el texto extraído para auditoría.
        
        Args:
            archivo: Archivo subido via FastAPI
            
        Returns:
            str: Texto extraído del archivo
            
        Raises:
            ValueError: Si el tipo de archivo no es soportado
        """
        if not archivo.filename:
            raise ValueError("Archivo sin nombre")
        
        extension = Path(archivo.filename).suffix.lower()
        contenido = await archivo.read()
        
        logger.info(f" Procesando archivo: {archivo.filename} ({extension})")
        
        # Determinar método de extracción según extensión
        if extension == '.pdf':
            texto = await self.extraer_texto_pdf(contenido, archivo.filename)
            # La función extraer_texto_pdf ya maneja automáticamente el OCR cuando es necesario
            return texto
        
        elif extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
            return await self.extraer_texto_imagen(contenido, archivo.filename)
        
        elif extension in ['.xlsx', '.xls']:
            return await self.extraer_texto_excel(contenido, archivo.filename)
        
        elif extension in ['.docx', '.doc']:
            return await self.extraer_texto_word(contenido, archivo.filename)
        
        elif extension in ['.msg', '.eml']:
            return await self.extraer_texto_emails(contenido, archivo.filename)

        elif extension == '.xml':
            return await self.extraer_texto_xml(contenido, archivo.filename)

        else:
            raise ValueError(f"Tipo de archivo no soportado: {extension}")
    
    def _evaluar_calidad_extraccion_pdf(self, texto_extraido: str, num_paginas: int) -> dict:
        """
        Evalúa la calidad del texto extraído de un PDF para determinar si se necesita OCR.
        
        Args:
            texto_extraido: Texto extraído del PDF
            num_paginas: Número total de páginas del PDF
            
        Returns:
            dict: Información sobre la calidad de extracción
        """
        # Contar páginas con mensaje de "vacía"
        mensajes_vacia = texto_extraido.count("[Página vacía o sin texto extraíble]")
        
        # Calcular texto útil (sin contar separadores y mensajes de páginas vacías)
        lineas = texto_extraido.split('\n')
        texto_util = ""
        
        for linea in lineas:
            # Excluir separadores de página y mensajes de páginas vacías
            if (not linea.startswith("--- PÁGINA") and 
                "[Página vacía o sin texto extraíble]" not in linea and
                linea.strip()):
                texto_util += linea + " "
        
        texto_util = texto_util.strip()
        caracteres_utiles = len(texto_util)
        
        # Calcular porcentajes
        porcentaje_paginas_vacias = (mensajes_vacia / num_paginas) * 100 if num_paginas > 0 else 0
        
        # Determinar si necesita OCR
        necesita_ocr = (
            porcentaje_paginas_vacias >= 80 or  # 80% o más páginas vacías
            caracteres_utiles < 100 or          # Menos de 100 caracteres útiles
            (porcentaje_paginas_vacias >= 50 and caracteres_utiles < 500)  # 50% vacías y poco texto
        )
        
        evaluacion = {
            "caracteres_totales": len(texto_extraido),
            "caracteres_utiles": caracteres_utiles,
            "paginas_totales": num_paginas,
            "paginas_vacias": mensajes_vacia,
            "porcentaje_paginas_vacias": porcentaje_paginas_vacias,
            "necesita_ocr": necesita_ocr,
            "razon_ocr": self._generar_razon_ocr(porcentaje_paginas_vacias, caracteres_utiles)
        }
        
        return evaluacion
    
    def _generar_razon_ocr(self, porcentaje_vacias: float, caracteres_utiles: int) -> str:
        """
        Genera una razón legible de por qué se necesita OCR.
        """
        if porcentaje_vacias >= 80:
            return f"80%+ páginas vacías ({porcentaje_vacias:.1f}%)"
        elif caracteres_utiles < 100:
            return f"Muy poco texto útil ({caracteres_utiles} caracteres)"
        elif porcentaje_vacias >= 50 and caracteres_utiles < 500:
            return f"50%+ páginas vacías ({porcentaje_vacias:.1f}%) y poco texto ({caracteres_utiles} caracteres)"
        else:
            return "Extracción satisfactoria"
    
    async def extraer_texto_pdf(self, contenido: bytes, nombre_archivo: str = "documento.pdf") -> str:
        """
        Extrae texto de archivo PDF usando PDF Plumber como método principal y PyPDF2 como fallback.
        GUARDA AUTOMÁTICAMENTE el texto extraído.
        
        Args:
            contenido: Contenido binario del archivo PDF
            nombre_archivo: Nombre del archivo original para guardado
            
        Returns:
            str: Texto extraído del PDF
        """
        # MÉTODO PRINCIPAL: PDF PLUMBER
        if PDFPLUMBER_DISPONIBLE:
            try:
                logger.info(f" Extrayendo texto con PDF Plumber (método principal): {nombre_archivo}")
                
                with pdfplumber.open(io.BytesIO(contenido)) as pdf:
                    texto_completo = ""
                    num_paginas = len(pdf.pages)
                    
                    logger.info(f" Procesando PDF con {num_paginas} página(s) usando PDF Plumber")
                    
                    for i, page in enumerate(pdf.pages):
                        # Extraer texto como fluye naturalmente
                        texto_pagina = page.extract_text()
                        if texto_pagina and texto_pagina.strip():  # Solo agregar si hay texto real
                            texto_completo += f"\n--- PÁGINA {i+1} ---\n{texto_pagina}\n"
                        else:
                            texto_completo += f"\n--- PÁGINA {i+1} ---\n[Página vacía o sin texto extraíble]\n"
                    
                    texto_final = texto_completo.strip()
                    
                    # EVALUAR CALIDAD DE EXTRACCIÓN
                    evaluacion = self._evaluar_calidad_extraccion_pdf(texto_final, num_paginas)
                    
                    # Preparar metadatos con evaluación
                    metadatos = {
                        "total_paginas": num_paginas,
                        "tamaño_archivo_bytes": len(contenido),
                        "metodo": "PDF Plumber (principal)",
                        "caracteres_extraidos": len(texto_final),
                        "evaluacion_calidad": evaluacion
                    }
                    
                    # SI NECESITA OCR, INTENTAR EXTRACCIÓN CON OCR INMEDIATAMENTE
                    if evaluacion["necesita_ocr"]:
                        logger.warning(f" PDF Plumber extrajo poco contenido útil: {evaluacion['razon_ocr']}")
                        logger.info(f" Intentando OCR automáticamente...")
                        
                        try:
                            texto_ocr = await self.extraer_texto_pdf_con_ocr(contenido, nombre_archivo)
                            
                            if texto_ocr and not texto_ocr.startswith("Error") and len(texto_ocr.strip()) > evaluacion["caracteres_utiles"]:
                                logger.info(f" OCR proporcionó mejor resultado que PDF Plumber")
                                logger.info(f" Comparación: PDF Plumber ({evaluacion['caracteres_utiles']} caracteres útiles) vs OCR ({len(texto_ocr.strip())} caracteres)")
                                return texto_ocr  # Retornar resultado de OCR
                            else:
                                logger.warning(f" OCR no mejoró el resultado, manteniendo extracción de PDF Plumber")
                                
                        except Exception as e:
                            logger.error(f" Error en OCR automático: {str(e)}")
                            logger.info(f" Continuando con resultado de PDF Plumber")
                    
                    # Guardar texto extraído automáticamente
                    archivo_guardado = self._guardar_texto_extraido(
                        nombre_archivo, texto_final, "PDF", metadatos
                    )
                    
                    logger.info(f" PDF Plumber: {len(texto_final)} caracteres extraídos")
                    return texto_final
                    
            except Exception as e:
                logger.warning(f"PDF Plumber falló: {str(e)}")
                logger.info(f"Intentando con PyPDF2 como fallback...")
        
        # MÉTODO FALLBACK: PyPDF2
        try:
            logger.info(f" Extrayendo texto con PyPDF2 (fallback): {nombre_archivo}")
            
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(contenido))
            texto_completo = ""
            
            num_paginas = len(pdf_reader.pages)
            logger.info(f" Procesando PDF con {num_paginas} página(s) usando PyPDF2")
            
            for i, page in enumerate(pdf_reader.pages):
                texto_pagina = page.extract_text()
                texto_completo += f"\n--- PÁGINA {i+1} ---\n{texto_pagina}\n"
            
            texto_final = texto_completo.strip()
            
            # Preparar metadatos
            metadatos = {
                "total_paginas": num_paginas,
                "tamaño_archivo_bytes": len(contenido),
                "metodo": "PyPDF2 (fallback)",
                "caracteres_extraidos": len(texto_final)
            }
            
            # Guardar texto extraído automáticamente
            archivo_guardado = self._guardar_texto_extraido(
                nombre_archivo, texto_final, "PDF", metadatos
            )
            
            logger.info(f"PyPDF2: {len(texto_final)} caracteres extraídos")
            return texto_final
            
        except Exception as e:
            error_msg = f"Error procesando PDF con ambos métodos (PDF Plumber + PyPDF2): {str(e)}"
            logger.error(f" {error_msg}")
            
            # Guardar también los errores para debugging
            self._guardar_texto_extraido(
                nombre_archivo, error_msg, "PDF_ERROR", {"error": str(e)}
            )
            
            return error_msg
    
    async def extraer_texto_pdf_con_ocr(self, contenido_pdf: bytes, nombre_archivo: str = "documento.pdf") -> str:
        """
        Extrae texto de PDF convirtiéndolo a imágenes y aplicando OCR.
        NUEVA FUNCIÓN para manejar PDFs con poco texto extraíble.
        
        Args:
            contenido_pdf: Contenido binario del archivo PDF
            nombre_archivo: Nombre del archivo original para guardado
            
        Returns:
            str: Texto extraído via OCR de las imágenes del PDF
        """
        try:
            # VALIDAR PDF ANTES DE INTENTAR CONVERSIÓN
            validacion = self._validar_pdf(contenido_pdf, nombre_archivo)
            
            if not validacion["valido"]:
                error_msg = f"PDF no válido para OCR: {validacion['error']}"
                logger.error(f" {error_msg}")
                
                # Guardar error detallado
                self._guardar_texto_extraido(
                    nombre_archivo, error_msg, "PDF_OCR_ERROR", 
                    {"error": validacion["error"], "validacion": validacion}
                )
                
                return error_msg
            
            logger.info(f" PDF validado: {validacion['info']['paginas']} páginas, {validacion['info']['tamaño_bytes']} bytes")
            logger.info(f" Método recomendado: {validacion['metodo_recomendado']}")
            
            # Convertir PDF a imágenes
            imagenes = self._convertir_pdf_a_imagenes(contenido_pdf, nombre_archivo)
            
            if not imagenes:
                error_msg = "No se pudieron convertir páginas del PDF a imágenes"
                logger.error(f" {error_msg}")
                
                self._guardar_texto_extraido(
                    nombre_archivo, error_msg, "PDF_OCR_ERROR", 
                    {"error": "Conversión PDF → Imagen falló", "validacion": validacion}
                )
                
                return error_msg
            
            # Aplicar OCR paralelo con ThreadPoolExecutor (2 workers fijos)
            texto_total, total_caracteres = await self._procesar_ocr_paralelo(imagenes, nombre_archivo)
            
            # Preparar metadatos
            metadatos = {
                "total_paginas": len(imagenes),
                "tamaño_archivo_bytes": len(contenido_pdf),
                "metodo": "PDF → Imagen → Google Vision OCR (Paralelo)",
                "workers_paralelos": 2,
                "caracteres_extraidos": total_caracteres,
                "paginas_procesadas": len(imagenes),
                "validacion": validacion["info"],
                "procesamiento_paralelo": True
            }
            
            # Guardar texto extraído automáticamente
            archivo_guardado = self._guardar_texto_extraido(
                nombre_archivo, texto_total, "PDF_OCR_PARALELO", metadatos
            )
            
            logger.info(f"OCR paralelo de PDF completado: {total_caracteres} caracteres de {len(imagenes)} paginas con 2 workers")
            
            return texto_total.strip()
            
        except Exception as e:
            error_msg = f"Error en OCR de PDF: {str(e)}"
            logger.error(f" {error_msg}")
            
            # Guardar error
            self._guardar_texto_extraido(
                nombre_archivo, error_msg, "PDF_OCR_ERROR", 
                {"error": str(e)}
            )
            
            return error_msg
    
    async def _procesar_ocr_paralelo(self, imagenes: list, nombre_archivo: str) -> tuple:
        """
        Procesa OCR de múltiples páginas en paralelo usando ThreadPoolExecutor.
        
        Args:
            imagenes: Lista de imágenes en bytes
            nombre_archivo: Nombre del archivo para logging
            
        Returns:
            tuple: (texto_total, total_caracteres)
        """
        if not imagenes:
            return "", 0
        
        # Configuración de OCR paralelo (2 workers fijos)
        max_workers = 2
        num_paginas = len(imagenes)
        
        # Logging específico sin emojis
        logger.info(f"Iniciando OCR paralelo: {num_paginas} paginas con {max_workers} workers")
        inicio_tiempo = asyncio.get_event_loop().time()
        
        # Función sincrónica para usar en ThreadPoolExecutor
        def aplicar_ocr_sincrono(imagen_bytes: bytes, num_pagina: int) -> tuple:
            """Función sincrónica que envuelve la llamada a Google Vision"""
            try:
                if not self.vision_client:
                    return num_pagina, "OCR no disponible - Google Vision no configurado"
                
                # Crear objeto Image para Vision
                image = vision.Image(content=imagen_bytes)
                
                # Detectar texto
                response = self.vision_client.text_detection(image=image)
                
                if response.error.message:
                    return num_pagina, f"Error en Vision API: {response.error.message}"
                
                texts = response.text_annotations
                
                if texts:
                    texto_extraido = texts[0].description
                    return num_pagina, texto_extraido
                else:
                    return num_pagina, ""
                    
            except Exception as e:
                return num_pagina, f"Error en OCR: {str(e)}"
        
        # Ejecutar OCR paralelo con ThreadPoolExecutor
        loop = asyncio.get_event_loop()
        
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Crear tareas para cada página
            tasks = [
                loop.run_in_executor(
                    executor, 
                    aplicar_ocr_sincrono, 
                    imagen_bytes, 
                    i + 1
                )
                for i, imagen_bytes in enumerate(imagenes)
            ]
            
            # Ejecutar todas las tareas en paralelo
            resultados = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Calcular tiempo transcurrido
        tiempo_total = asyncio.get_event_loop().time() - inicio_tiempo
        
        # Procesar resultados manteniendo orden de páginas
        texto_total = ""
        total_caracteres = 0
        paginas_exitosas = 0
        
        for resultado in resultados:
            if isinstance(resultado, Exception):
                # Manejar excepción
                num_pagina = len([r for r in resultados[:resultados.index(resultado)] if not isinstance(r, Exception)]) + 1
                texto_total += f"\n--- PÁGINA {num_pagina} (OCR) ---\n[Error en procesamiento paralelo: {resultado}]\n"
                continue
            
            num_pagina, texto_pagina = resultado
            
            if texto_pagina and not texto_pagina.startswith("Error"):
                texto_total += f"\n--- PÁGINA {num_pagina} (OCR) ---\n{texto_pagina}\n"
                total_caracteres += len(texto_pagina)
                paginas_exitosas += 1
            else:
                texto_total += f"\n--- PÁGINA {num_pagina} (OCR) ---\n[Error en OCR o página vacía]\n"
        
        # Logging de resultados sin emojis
        logger.info(f"OCR paralelo completado: {paginas_exitosas}/{num_paginas} paginas exitosas")
        logger.info(f"Tiempo total de OCR paralelo: {tiempo_total:.2f} segundos")
        logger.info(f"Promedio por pagina: {tiempo_total/num_paginas:.2f} segundos")
        logger.info(f"Caracteres extraidos: {total_caracteres}")
        
        return texto_total, total_caracteres
    
    async def _aplicar_ocr_a_imagen(self, imagen_bytes: bytes, descripcion: str = "imagen") -> str:
        """
        Aplica OCR a una imagen específica de  manera directa.
        
        Args:
            imagen_bytes: Bytes de la imagen
            descripcion: Descripción para logging
            
        Returns:
            str: Texto extraído de la imagen
        """
        if not self.vision_client:
            return "OCR no disponible - Google Vision no configurado"
        
        try:
            # Crear objeto Image para Vision
            image = vision.Image(content=imagen_bytes)
            
            # Detectar texto
            response = self.vision_client.text_detection(image=image)
            
            if response.error.message:
                raise Exception(f'Error en Vision API: {response.error.message}')
            
            texts = response.text_annotations
            
            if texts:
                texto_extraido = texts[0].description
                logger.debug(f" OCR exitoso en {descripcion}: {len(texto_extraido)} caracteres")
                return texto_extraido
            else:
                logger.debug(f" No se detectó texto en {descripcion}")
                return ""
                
        except Exception as e:
            logger.error(f" Error en OCR de {descripcion}: {e}")
            return f"Error en OCR: {str(e)}"
    
    async def extraer_texto_imagen(self, contenido: bytes, nombre_archivo: str = "imagen.jpg", metodo: str = "OCR") -> str:
        """
        Extrae texto de imagen usando Google Vision OCR.
        GUARDA AUTOMÁTICAMENTE el texto extraído.
        
        Args:
            contenido: Contenido binario de la imagen
            nombre_archivo: Nombre del archivo original para guardado
            metodo: Método específico de OCR (para casos como fallback)
            
        Returns:
            str: Texto extraído de la imagen
        """
        if not self.vision_client:
            error_msg = "OCR no disponible - Google Vision no configurado"
            logger.warning(f" {error_msg}")
            
            # Guardar el error
            self._guardar_texto_extraido(
                nombre_archivo, error_msg, f"{metodo}_ERROR", 
                {"error": "Google Vision no configurado"}
            )
            
            return error_msg
        
        try:
            # Aplicar OCR
            texto_extraido = await self._aplicar_ocr_a_imagen(contenido, nombre_archivo)
            
            if texto_extraido and not texto_extraido.startswith("Error"):
                # Preparar metadatos
                metadatos = {
                    "tamaño_archivo_bytes": len(contenido),
                    "metodo": "Google Vision OCR",
                    "caracteres_extraidos": len(texto_extraido)
                }
                
                # Guardar texto extraído automáticamente
                archivo_guardado = self._guardar_texto_extraido(
                    nombre_archivo, texto_extraido, metodo, metadatos
                )
                
                logger.info(f" OCR exitoso: {len(texto_extraido)} caracteres extraídos")
                return texto_extraido
            else:
                no_texto_msg = "No se detectó texto en la imagen"
                logger.warning(f" {no_texto_msg}")
                
                # Guardar resultado vacío
                self._guardar_texto_extraido(
                    nombre_archivo, no_texto_msg, f"{metodo}_VACIO", 
                    {"elementos_detectados": 0}
                )
                
                return no_texto_msg
                
        except Exception as e:
            error_msg = f"Error en OCR con Google Vision: {str(e)}"
            logger.error(f" {error_msg}")
            
            # Guardar error
            self._guardar_texto_extraido(
                nombre_archivo, error_msg, f"{metodo}_ERROR", 
                {"error": str(e)}
            )
            
            return error_msg
    
    async def extraer_texto_excel(self, contenido: bytes, nombre_archivo: str = "archivo.xlsx") -> str:
        """
        Extrae texto de archivo Excel (XLSX/XLS).
        GUARDA AUTOMÁTICAMENTE el texto extraído.
        
        Args:
            contenido: Contenido binario del archivo Excel
            nombre_archivo: Nombre del archivo original para guardado
            
        Returns:
            str: Texto extraído y formateado del Excel
        """
        try:
            # Leer Excel
            df = pd.read_excel(io.BytesIO(contenido), sheet_name=None)  # Leer todas las hojas
            
            texto_completo = ""
            total_hojas = 0
            total_filas = 0
            
            # Si hay múltiples hojas
            if isinstance(df, dict):
                total_hojas = len(df)
                for nombre_hoja, dataframe in df.items():
                    texto_completo += f"\n--- HOJA: {nombre_hoja} ---\n"
                    texto_completo += dataframe.to_string(index=False, na_rep='')
                    texto_completo += "\n"
                    total_filas += len(dataframe)
            else:
                # Una sola hoja
                total_hojas = 1
                total_filas = len(df)
                texto_completo = df.to_string(index=False, na_rep='')
            
            texto_final = texto_completo.strip()
            
            # Preparar metadatos
            metadatos = {
                "total_hojas": total_hojas,
                "total_filas": total_filas,
                "tamaño_archivo_bytes": len(contenido),
                "metodo": "pandas.read_excel"
            }
            
            # Guardar texto extraído automáticamente
            archivo_guardado = self._guardar_texto_extraido(
                nombre_archivo, texto_final, "EXCEL", metadatos
            )
            
            logger.info(f" Excel procesado: {len(texto_final)} caracteres extraídos")
            logger.info(f" Hojas: {total_hojas}, Filas: {total_filas}")
            
            return texto_final
            
        except Exception as e:
            error_msg = f"Error procesando Excel: {str(e)}"
            logger.error(f" {error_msg}")
            
            # Guardar error
            self._guardar_texto_extraido(
                nombre_archivo, error_msg, "EXCEL_ERROR", 
                {"error": str(e)}
            )
            
            return error_msg
    
    async def extraer_texto_word(self, contenido: bytes, nombre_archivo: str = "documento.docx") -> str:
        """
        Extrae texto de archivo Word (DOCX).
        GUARDA AUTOMÁTICAMENTE el texto extraído.
        
        Args:
            contenido: Contenido binario del archivo Word
            nombre_archivo: Nombre del archivo original para guardado
            
        Returns:
            str: Texto extraído del documento Word
        """
        try:
            # Leer documento Word
            doc = Document(io.BytesIO(contenido))
            texto_completo = ""
            total_parrafos = 0
            total_tablas = 0
            
            # Extraer párrafos
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Solo párrafos con contenido
                    texto_completo += paragraph.text + "\n"
                    total_parrafos += 1
            
            # Extraer tablas si las hay
            if doc.tables:
                total_tablas = len(doc.tables)
                texto_completo += "\n--- TABLAS DEL DOCUMENTO ---\n"
                for i, table in enumerate(doc.tables):
                    texto_completo += f"\nTabla {i+1}:\n"
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text.strip():
                            texto_completo += row_text + "\n"
            
            texto_final = texto_completo.strip()
            
            # Preparar metadatos
            metadatos = {
                "total_parrafos": total_parrafos,
                "total_tablas": total_tablas,
                "tamaño_archivo_bytes": len(contenido),
                "metodo": "python-docx"
            }
            
            # Guardar texto extraído automáticamente
            archivo_guardado = self._guardar_texto_extraido(
                nombre_archivo, texto_final, "WORD", metadatos
            )
            
            logger.info(f" Word procesado: {len(texto_final)} caracteres extraídos")
            logger.info(f" Párrafos: {total_parrafos}, Tablas: {total_tablas}")
            
            return texto_final
            
        except Exception as e:
            error_msg = f"Error procesando Word: {str(e)}"
            logger.error(f" {error_msg}")
            
            # Guardar error
            self._guardar_texto_extraido(
                nombre_archivo, error_msg, "WORD_ERROR", 
                {"error": str(e)}
            )
            
            return error_msg
    
    async def extraer_texto_emails(self, contenido: bytes, nombre_archivo: str = "email") -> str:
        """
        Extrae texto y metadatos de archivos de email (.msg y .eml).
        GUARDA AUTOMÁTICAMENTE el texto extraído con formato estructurado.
        
        Args:
            contenido: Contenido binario del archivo de email
            nombre_archivo: Nombre del archivo original para guardado
            
        Returns:
            str: Texto extraído con metadatos del email formateado
        """
        try:
            extension = Path(nombre_archivo).suffix.lower()
            
            if extension == '.msg':
                return await self._procesar_msg(contenido, nombre_archivo)
            elif extension == '.eml':
                return await self._procesar_eml(contenido, nombre_archivo)
            else:
                error_msg = f"Extensión de email no soportada: {extension}"
                logger.error(f" {error_msg}")
                
                # Guardar error
                self._guardar_texto_extraido(
                    nombre_archivo, error_msg, "EMAIL_ERROR", 
                    {"error": "Extensión no soportada", "extension": extension}
                )
                
                return error_msg
                
        except Exception as e:
            error_msg = f"Error procesando email: {str(e)}"
            logger.error(f" {error_msg}")
            
            # Guardar error
            self._guardar_texto_extraido(
                nombre_archivo, error_msg, "EMAIL_ERROR", 
                {"error": str(e)}
            )
            
            return error_msg
    
    async def _procesar_msg(self, contenido: bytes, nombre_archivo: str) -> str:
        """
        Procesa archivos .msg usando extract-msg.
        
        Args:
            contenido: Contenido binario del archivo .msg
            nombre_archivo: Nombre del archivo para logging
            
        Returns:
            str: Texto extraído formateado del email
        """
        if not EXTRACT_MSG_DISPONIBLE:
            error_msg = "Librería extract-msg no disponible. Instale con: pip install extract-msg"
            logger.error(f" {error_msg}")
            
            self._guardar_texto_extraido(
                nombre_archivo, error_msg, "MSG_ERROR", 
                {"error": "extract-msg no instalado"}
            )
            
            return error_msg
        
        try:
            # Crear archivo temporal para extract-msg
            import tempfile
            with tempfile.NamedTemporaryFile(suffix='.msg', delete=False) as temp_file:
                temp_file.write(contenido)
                temp_path = temp_file.name
            
            try:
                # Extraer con extract-msg
                msg = extract_msg.Message(temp_path)
                
                # Extraer metadatos
                asunto = msg.subject or "[Sin asunto]"
                remitente = self._formatear_direccion(msg.sender)
                destinatarios = self._extraer_destinatarios_msg(msg)
                fecha = self._formatear_fecha(msg.date)
                cuerpo = self._extraer_cuerpo_msg(msg)
                adjuntos = self._listar_adjuntos_msg(msg)
                
                # Formatear texto final
                texto_formateado = self._formatear_email(
                    asunto, remitente, destinatarios, fecha, cuerpo, adjuntos, "MSG"
                )
                
                # Preparar metadatos
                metadatos = {
                    "tipo_archivo": "MSG",
                    "asunto": asunto,
                    "remitente": remitente,
                    "destinatarios": destinatarios,
                    "fecha": fecha,
                    "adjuntos_detectados": len(adjuntos),
                    "tamaño_archivo_bytes": len(contenido),
                    "metodo": "extract-msg"
                }
                
                # Guardar texto extraído
                self._guardar_texto_extraido(
                    nombre_archivo, texto_formateado, "EMAIL_MSG", metadatos
                )
                
                logger.info(f" Email .msg procesado: {len(texto_formateado)} caracteres")
                logger.info(f" Asunto: {asunto[:50]}...")
                
                return texto_formateado
                
            finally:
                # Limpiar archivo temporal
                import os
                try:
                    os.unlink(temp_path)
                except:
                    pass
                    
        except Exception as e:
            error_msg = f"Error procesando archivo .msg: {str(e)}"
            logger.error(f" {error_msg}")
            
            self._guardar_texto_extraido(
                nombre_archivo, error_msg, "MSG_ERROR", 
                {"error": str(e)}
            )
            
            return error_msg
    
    async def _procesar_eml(self, contenido: bytes, nombre_archivo: str) -> str:
        """
        Procesa archivos .eml usando la librería email estándar.
        
        Args:
            contenido: Contenido binario del archivo .eml
            nombre_archivo: Nombre del archivo para logging
            
        Returns:
            str: Texto extraído formateado del email
        """
        try:
            # Decodificar contenido a string
            contenido_str = self._decodificar_email(contenido)
            
            # Parsear email
            msg = email.message_from_string(contenido_str)
            
            # Extraer metadatos
            asunto = self._decodificar_header(msg.get('Subject', '[Sin asunto]'))
            remitente = self._decodificar_header(msg.get('From', '[Remitente desconocido]'))
            destinatarios = self._extraer_destinatarios_eml(msg)
            fecha = self._formatear_fecha_eml(msg.get('Date'))
            cuerpo = self._extraer_cuerpo_eml(msg)
            adjuntos = self._listar_adjuntos_eml(msg)
            
            # Formatear texto final
            texto_formateado = self._formatear_email(
                asunto, remitente, destinatarios, fecha, cuerpo, adjuntos, "EML"
            )
            
            # Preparar metadatos
            metadatos = {
                "tipo_archivo": "EML",
                "asunto": asunto,
                "remitente": remitente,
                "destinatarios": destinatarios,
                "fecha": fecha,
                "adjuntos_detectados": len(adjuntos),
                "tamaño_archivo_bytes": len(contenido),
                "metodo": "email estándar"
            }
            
            # Guardar texto extraído
            self._guardar_texto_extraido(
                nombre_archivo, texto_formateado, "EMAIL_EML", metadatos
            )
            
            logger.info(f" Email .eml procesado: {len(texto_formateado)} caracteres")
            logger.info(f" Asunto: {asunto[:50]}...")
            
            return texto_formateado
            
        except Exception as e:
            error_msg = f"Error procesando archivo .eml: {str(e)}"
            logger.error(f"{error_msg}")
            
            self._guardar_texto_extraido(
                nombre_archivo, error_msg, "EML_ERROR", 
                {"error": str(e)}
            )
            
            return error_msg
    
    def _decodificar_email(self, contenido: bytes) -> str:
        """
        Decodifica el contenido de un email manejando diferentes codificaciones.
        """
        codificaciones = ['utf-8', 'latin1', 'cp1252', 'iso-8859-1']
        
        for codificacion in codificaciones:
            try:
                return contenido.decode(codificacion)
            except UnicodeDecodeError:
                continue
        
        # Fallback con errores ignorados
        return contenido.decode('utf-8', errors='ignore')
    
    def _decodificar_header(self, header_str: str) -> str:
        """
        Decodifica headers de email que pueden estar codificados.
        """
        if not header_str:
            return ""
        
        try:
            decoded_fragments = decode_header(header_str)
            decoded_string = ""
            
            for fragment, encoding in decoded_fragments:
                if isinstance(fragment, bytes):
                    if encoding:
                        try:
                            decoded_string += fragment.decode(encoding)
                        except:
                            decoded_string += fragment.decode('utf-8', errors='ignore')
                    else:
                        decoded_string += fragment.decode('utf-8', errors='ignore')
                else:
                    decoded_string += fragment
            
            return decoded_string.strip()
            
        except Exception:
            return header_str
    
    def _formatear_direccion(self, direccion: str) -> str:
        """
        Formatea una dirección de email para mostrar nombre y email.
        """
        if not direccion:
            return "[Desconocido]"
        
        try:
            nombre, email_addr = parseaddr(direccion)
            if nombre and email_addr:
                return f"{nombre} <{email_addr}>"
            elif email_addr:
                return email_addr
            else:
                return direccion
        except:
            return direccion
    
    def _extraer_destinatarios_msg(self, msg) -> str:
        """
        Extrae destinatarios de un mensaje .msg.
        """
        destinatarios = []
        
        # To
        if hasattr(msg, 'to') and msg.to:
            destinatarios.append(f"Para: {msg.to}")
        
        # CC
        if hasattr(msg, 'cc') and msg.cc:
            destinatarios.append(f"CC: {msg.cc}")
        
        # BCC
        if hasattr(msg, 'bcc') and msg.bcc:
            destinatarios.append(f"BCC: {msg.bcc}")
        
        return "; ".join(destinatarios) if destinatarios else "[Sin destinatarios]"
    
    def _extraer_destinatarios_eml(self, msg) -> str:
        """
        Extrae destinatarios de un mensaje .eml.
        """
        destinatarios = []
        
        # To
        to_header = msg.get('To')
        if to_header:
            destinatarios.append(f"Para: {self._decodificar_header(to_header)}")
        
        # CC
        cc_header = msg.get('Cc')
        if cc_header:
            destinatarios.append(f"CC: {self._decodificar_header(cc_header)}")
        
        # BCC
        bcc_header = msg.get('Bcc')
        if bcc_header:
            destinatarios.append(f"BCC: {self._decodificar_header(bcc_header)}")
        
        return "; ".join(destinatarios) if destinatarios else "[Sin destinatarios]"
    
    def _formatear_fecha(self, fecha) -> str:
        """
        Formatea fecha de mensaje .msg.
        """
        if not fecha:
            return "[Fecha desconocida]"
        
        try:
            if hasattr(fecha, 'strftime'):
                return fecha.strftime("%Y-%m-%d %H:%M:%S")
            else:
                return str(fecha)
        except:
            return str(fecha)
    
    def _formatear_fecha_eml(self, fecha_str: str) -> str:
        """
        Formatea fecha de mensaje .eml.
        """
        if not fecha_str:
            return "[Fecha desconocida]"
        
        try:
            fecha_obj = parsedate_to_datetime(fecha_str)
            return fecha_obj.strftime("%Y-%m-%d %H:%M:%S")
        except:
            return fecha_str
    
    def _extraer_cuerpo_msg(self, msg) -> str:
        """
        Extrae el cuerpo del mensaje .msg.
        """
        cuerpo = ""
        
        try:
            # Intentar texto plano primero
            if hasattr(msg, 'body') and msg.body:
                cuerpo = msg.body
            # Fallback a HTML
            elif hasattr(msg, 'htmlBody') and msg.htmlBody:
                cuerpo = self._html_a_texto(msg.htmlBody)
                cuerpo = f"[CONVERTIDO DE HTML]\n{cuerpo}"
            else:
                cuerpo = "[Sin contenido de mensaje]"
                
        except Exception as e:
            cuerpo = f"[Error extrayendo cuerpo: {str(e)}]"
        
        return cuerpo.strip() if cuerpo else "[Mensaje vacío]"
    
    def _extraer_cuerpo_eml(self, msg) -> str:
        """
        Extrae el cuerpo del mensaje .eml.
        """
        cuerpo = ""
        
        try:
            if msg.is_multipart():
                # Buscar partes de texto
                for part in msg.walk():
                    content_type = part.get_content_type()
                    
                    if content_type == "text/plain":
                        payload = part.get_payload(decode=True)
                        if payload:
                            charset = part.get_content_charset() or 'utf-8'
                            try:
                                cuerpo = payload.decode(charset)
                                break  # Priorizar texto plano
                            except:
                                cuerpo = payload.decode('utf-8', errors='ignore')
                                break
                    
                    elif content_type == "text/html" and not cuerpo:
                        payload = part.get_payload(decode=True)
                        if payload:
                            charset = part.get_content_charset() or 'utf-8'
                            try:
                                html_content = payload.decode(charset)
                                cuerpo = self._html_a_texto(html_content)
                                cuerpo = f"[CONVERTIDO DE HTML]\n{cuerpo}"
                            except:
                                html_content = payload.decode('utf-8', errors='ignore')
                                cuerpo = self._html_a_texto(html_content)
                                cuerpo = f"[CONVERTIDO DE HTML]\n{cuerpo}"
            else:
                # Mensaje simple
                payload = msg.get_payload(decode=True)
                if payload:
                    charset = msg.get_content_charset() or 'utf-8'
                    try:
                        cuerpo = payload.decode(charset)
                    except:
                        cuerpo = payload.decode('utf-8', errors='ignore')
                        
        except Exception as e:
            cuerpo = f"[Error extrayendo cuerpo: {str(e)}]"
        
        return cuerpo.strip() if cuerpo else "[Mensaje vacío]"
    
    def _html_a_texto(self, html_content: str) -> str:
        """
        Convierte contenido HTML a texto plano simple.
        """
        try:
            import re
            
            # Remover scripts y styles
            html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
            html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
            
            # Reemplazar saltos de línea HTML
            html_content = re.sub(r'<br[^>]*>', '\n', html_content, flags=re.IGNORECASE)
            html_content = re.sub(r'</p>', '\n\n', html_content, flags=re.IGNORECASE)
            html_content = re.sub(r'</div>', '\n', html_content, flags=re.IGNORECASE)
            
            # Remover todas las etiquetas HTML
            texto_limpio = re.sub(r'<[^>]+>', '', html_content)
            
            # Limpiar espacios múltiples y líneas vacías
            texto_limpio = re.sub(r'\n\s*\n', '\n\n', texto_limpio)
            texto_limpio = re.sub(r' +', ' ', texto_limpio)
            
            return texto_limpio.strip()
            
        except Exception:
            # Fallback simple
            import re
            return re.sub(r'<[^>]+>', '', html_content)
    
    def _listar_adjuntos_msg(self, msg) -> list:
        """
        Lista los adjuntos de un mensaje .msg.
        """
        adjuntos = []
        
        try:
            if hasattr(msg, 'attachments') and msg.attachments:
                for attachment in msg.attachments:
                    try:
                        nombre = getattr(attachment, 'longFilename', None) or getattr(attachment, 'shortFilename', 'adjunto_sin_nombre')
                        tamaño = getattr(attachment, 'size', 0)
                        adjuntos.append(f"{nombre} ({tamaño} bytes)")
                    except:
                        adjuntos.append("adjunto_sin_info")
        except Exception:
            pass
        
        return adjuntos
    
    def _listar_adjuntos_eml(self, msg) -> list:
        """
        Lista los adjuntos de un mensaje .eml.
        """
        adjuntos = []
        
        try:
            for part in msg.walk():
                content_disposition = part.get('Content-Disposition', '')
                
                if content_disposition and 'attachment' in content_disposition:
                    filename = part.get_filename()
                    if filename:
                        filename = self._decodificar_header(filename)
                        tamaño = len(part.get_payload(decode=True) or b'')
                        adjuntos.append(f"{filename} ({tamaño} bytes)")
                    else:
                        adjuntos.append("adjunto_sin_nombre")
        except Exception:
            pass
        
        return adjuntos
    
    def _formatear_email(self, asunto: str, remitente: str, destinatarios: str, 
                        fecha: str, cuerpo: str, adjuntos: list, tipo: str) -> str:
        """
        Formatea toda la información del email en un texto estructurado.
        """
        separador = "=" * 60
        
        texto_formateado = f"""=== INFORMACIÓN DEL EMAIL ({tipo}) ===
ASUNTO: {asunto}
REMITENTE: {remitente}
DESTINATARIOS: {destinatarios}
FECHA: {fecha}

{separador}
=== CUERPO DEL EMAIL ===
{separador}

{cuerpo}

{separador}
=== ARCHIVOS ADJUNTOS ===
{separador}
"""
        
        if adjuntos:
            for i, adjunto in enumerate(adjuntos, 1):
                texto_formateado += f"\n{i}. {adjunto}"
        else:
            texto_formateado += "\n[Sin archivos adjuntos]"
        
        texto_formateado += f"\n\n{separador}\n=== FIN DEL EMAIL ===\n{separador}"

        return texto_formateado

    async def extraer_texto_xml(self, contenido: bytes, nombre_archivo: str = "documento.xml") -> str:
        """
        Extrae texto estructurado de archivo XML preservando jerarquía.
        Optimizado para facturas electrónicas colombianas (UBL 2.1).

        Características:
        - Omite firmas digitales (ds:Signature, xades:*)
        - Extrae CDATA como texto crudo sin parsear
        - Sin priorización: extrae toda la estructura
        - Formato simple con indentación

        Args:
            contenido: Contenido binario del archivo XML
            nombre_archivo: Nombre del archivo original

        Returns:
            str: Texto estructurado del XML
        """
        try:
            # 1. Detectar encoding
            encoding = self._detectar_encoding_xml(contenido)

            # 2. Parsear XML
            tree = ET.ElementTree(ET.fromstring(contenido))
            root = tree.getroot()

            # 3. Extraer namespaces
            namespaces = self._extraer_namespaces(root)

            # 4. Detectar tipo de documento
            tipo_documento = self._detectar_tipo_documento_xml(root, namespaces)

            # 5. Extraer jerarquía (omitiendo firmas)
            texto_completo = f"=== {tipo_documento} ===\n\n"
            texto_completo += f"Elemento Raíz: {self._limpiar_nombre_elemento(root.tag)}\n\n"

            estadisticas = {"total_elementos": 0, "profundidad_maxima": 0, "firmas_omitidas": 0}

            texto_jerarquia = self._extraer_jerarquia_xml(
                root,
                nivel=0,
                estadisticas=estadisticas,
                omitir_firmas=True
            )

            texto_completo += texto_jerarquia

            # 6. Metadatos
            metadatos = {
                "tipo_documento": tipo_documento,
                "elemento_raiz": self._limpiar_nombre_elemento(root.tag),
                "namespaces_detectados": namespaces,
                "total_elementos": estadisticas["total_elementos"],
                "profundidad_maxima": estadisticas["profundidad_maxima"],
                "firmas_omitidas": estadisticas["firmas_omitidas"],
                "tamaño_archivo_bytes": len(contenido),
                "encoding": encoding,
                "metodo": "xml.etree.ElementTree"
            }

            # 7. Guardar automáticamente
            self._guardar_texto_extraido(nombre_archivo, texto_completo, "XML", metadatos)

            logger.info(f"XML procesado: {len(texto_completo)} caracteres")
            logger.info(f"Elementos: {estadisticas['total_elementos']}, Firmas omitidas: {estadisticas['firmas_omitidas']}")

            return texto_completo

        except ET.ParseError as e:
            error_msg = f"Error parseando XML: {str(e)}"
            self._guardar_texto_extraido(nombre_archivo, error_msg, "XML_ERROR", {"error": str(e)})
            return error_msg
        except Exception as e:
            error_msg = f"Error procesando XML: {str(e)}"
            self._guardar_texto_extraido(nombre_archivo, error_msg, "XML_ERROR", {"error": str(e)})
            return error_msg

    def _detectar_encoding_xml(self, contenido: bytes) -> str:
        """
        Detecta encoding del XML desde declaración o BOM.

        Args:
            contenido: Contenido binario del archivo XML

        Returns:
            str: Encoding detectado
        """
        # Buscar <?xml encoding="..."?>
        if b'<?xml' in contenido[:200]:
            declaracion = contenido[:200].split(b'?>')[0]
            if b'encoding=' in declaracion:
                # Extraer encoding entre comillas
                import re
                match = re.search(rb'encoding=["\']([^"\']+)["\']', declaracion)
                if match:
                    return match.group(1).decode('ascii').lower()

        # UTF-8 BOM
        if contenido.startswith(b'\xef\xbb\xbf'):
            return 'utf-8-sig'

        return 'utf-8'

    def _extraer_namespaces(self, root) -> dict:
        """
        Extrae namespaces del documento XML.

        Args:
            root: Elemento raíz del XML

        Returns:
            dict: Diccionario de namespaces
        """
        namespaces = {}

        # Extraer namespace del tag raíz
        if '}' in root.tag:
            default_ns = root.tag.split('}')[0][1:]
            namespaces['default'] = default_ns

        # Intentar extraer de atributos (puede no estar disponible en ElementTree)
        for key, value in root.attrib.items():
            if key == 'xmlns':
                namespaces['default'] = value
            elif key.startswith('xmlns:'):
                prefix = key.split(':', 1)[1]
                namespaces[prefix] = value

        # Si no se encontró namespace por atributos, usar el del tag
        if not namespaces and '}' in root.tag:
            default_ns = root.tag.split('}')[0][1:]
            namespaces['default'] = default_ns

        return namespaces

    def _detectar_tipo_documento_xml(self, root, namespaces: dict) -> str:
        """
        Detecta tipo de documento basándose en raíz y namespaces.

        Args:
            root: Elemento raíz del XML
            namespaces: Namespaces del documento

        Returns:
            str: Tipo de documento detectado
        """
        nombre_raiz = self._limpiar_nombre_elemento(root.tag)

        if 'oasis' in str(namespaces.get('default', '')):
            if 'Invoice' in nombre_raiz:
                return "Factura Electrónica UBL 2.1"
            elif 'AttachedDocument' in nombre_raiz:
                return "Documento Adjunto de Factura Electrónica"
            elif 'CreditNote' in nombre_raiz:
                return "Nota Crédito Electrónica UBL 2.1"
            else:
                return f"Documento UBL 2.1 ({nombre_raiz})"

        return f"Documento XML ({nombre_raiz})"

    def _limpiar_nombre_elemento(self, tag: str) -> str:
        """
        Remueve namespace URI del tag: {uri}Element → Element.

        Args:
            tag: Tag XML con o sin namespace

        Returns:
            str: Nombre del elemento limpio
        """
        if '}' in tag:
            return tag.split('}', 1)[1]
        return tag

    def _extraer_jerarquia_xml(self, elemento, nivel: int, estadisticas: dict,
                               omitir_firmas: bool = True) -> str:
        """
        Extrae jerarquía XML recursivamente.

        Características:
        - Omite firmas digitales si omitir_firmas=True
        - Extrae CDATA como texto crudo
        - Formato simple con indentación

        Args:
            elemento: Elemento XML a procesar
            nivel: Nivel de profundidad actual
            estadisticas: Dict para rastrear estadísticas
            omitir_firmas: Si debe omitir firmas digitales

        Returns:
            str: Texto jerárquico del XML
        """
        texto = ""
        indentacion = "  " * nivel

        # Obtener nombre limpio
        nombre = self._limpiar_nombre_elemento(elemento.tag)

        # OMITIR FIRMAS DIGITALES
        if omitir_firmas:
            if nombre == 'Signature' or nombre.startswith('Qualifying') or nombre.startswith('Signed'):
                estadisticas["firmas_omitidas"] += 1
                return indentacion + "[FIRMA DIGITAL OMITIDA]\n"

            # Omitir UBLExtension que solo contiene firmas
            if nombre == 'UBLExtension':
                # Verificar si contiene solo ExtensionContent > Signature
                for hijo in list(elemento):
                    hijo_nombre = self._limpiar_nombre_elemento(hijo.tag)
                    if hijo_nombre == 'ExtensionContent':
                        for nieto in list(hijo):
                            nieto_nombre = self._limpiar_nombre_elemento(nieto.tag)
                            if nieto_nombre == 'Signature':
                                estadisticas["firmas_omitidas"] += 1
                                return indentacion + "[FIRMA DIGITAL OMITIDA]\n"

        # Actualizar estadísticas
        estadisticas["total_elementos"] += 1
        if nivel > estadisticas["profundidad_maxima"]:
            estadisticas["profundidad_maxima"] = nivel

        # Línea del elemento
        linea = f"{indentacion}{nombre}"

        # Agregar atributos importantes (excluyendo xmlns)
        atributos = []
        for key, value in elemento.attrib.items():
            if not key.startswith('{') and not key.startswith('xmlns'):
                key_limpio = self._limpiar_nombre_elemento(key)
                atributos.append(f"{key_limpio}='{value}'")

        if atributos:
            linea += f" [{', '.join(atributos)}]"

        # Texto del elemento
        texto_elemento = (elemento.text or "").strip()

        # MANEJO DE CDATA
        if texto_elemento:
            # Detectar si es CDATA con XML (comienza con <?xml)
            if texto_elemento.startswith('<?xml'):
                texto += linea + "\n"
                texto += indentacion + "  [CDATA - Factura XML Embebida - INICIO]\n"
                # Mostrar TODAS las líneas del CDATA sin truncar
                lineas_cdata = texto_elemento.split('\n')
                for linea_cdata in lineas_cdata:
                    texto += indentacion + "  " + linea_cdata + "\n"
                texto += indentacion + "  [CDATA - Factura XML Embebida - FIN]\n"
            else:
                # Texto normal
                if len(texto_elemento) > 200:
                    texto_elemento = texto_elemento[:200] + "..."
                texto_elemento = texto_elemento.replace('\n', ' ')
                linea += f": {texto_elemento}"
                texto += linea + "\n"
        else:
            texto += linea + "\n"

        # Procesar hijos recursivamente
        hijos = list(elemento)
        if hijos:
            for hijo in hijos:
                texto += self._extraer_jerarquia_xml(hijo, nivel + 1, estadisticas, omitir_firmas)

        return texto

    def validar_archivo(self, archivo: UploadFile) -> Dict[str, Any]:
        """
        Valida si un archivo es procesable y retorna información sobre él.
        
        Args:
            archivo: Archivo a validar
            
        Returns:
            Dict con información de validación
        """
        extensiones_soportadas = ['.pdf', '.jpg', '.jpeg', '.png', '.gif', '.bmp',
                                '.tiff', '.xlsx', '.xls', '.docx', '.doc', '.msg', '.eml', '.xml']
        
        if not archivo.filename:
            return {
                "valido": False,
                "error": "Archivo sin nombre"
            }
        
        extension = Path(archivo.filename).suffix.lower()
        
        if extension not in extensiones_soportadas:
            return {
                "valido": False,
                "error": f"Extensión no soportada: {extension}",
                "extensiones_soportadas": extensiones_soportadas
            }
        
        # Determinar tipo de procesamiento
        tipo_procesamiento = "Desconocido"
        guardado_automatico = True
        
        if extension == '.pdf':
            if PDFPLUMBER_DISPONIBLE:
                tipo_procesamiento = "PDF Plumber (principal) + PyPDF2 (fallback)"
                if PDF2IMAGE_DISPONIBLE or PYMUPDF_DISPONIBLE:
                    tipo_procesamiento += " + OCR con conversión a imagen"
            else:
                tipo_procesamiento = "PyPDF2 (sin PDF Plumber)"
                if PDF2IMAGE_DISPONIBLE or PYMUPDF_DISPONIBLE:
                    tipo_procesamiento += " + OCR con conversión a imagen"
                else:
                    tipo_procesamiento += " (OCR fallback limitado)"
        elif extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
            tipo_procesamiento = "OCR con Google Vision"
        elif extension in ['.xlsx', '.xls']:
            tipo_procesamiento = "Procesamiento Excel"
        elif extension in ['.docx', '.doc']:
            tipo_procesamiento = "Procesamiento Word"
        elif extension in ['.msg', '.eml']:
            if extension == '.msg' and EXTRACT_MSG_DISPONIBLE:
                tipo_procesamiento = "Procesamiento Email (.msg) con extract-msg"
            elif extension == '.eml':
                tipo_procesamiento = "Procesamiento Email (.eml) con email estándar"
            else:
                tipo_procesamiento = "Procesamiento Email (dependencias limitadas)"
        elif extension == '.xml':
            tipo_procesamiento = "Procesamiento XML con xml.etree.ElementTree (facturas electrónicas)"


        return {
            "valido": True,
            "extension": extension,
            "tipo_procesamiento": tipo_procesamiento,
            "ocr_disponible": self.vision_client is not None,
            "pdf_to_image_disponible": PDF2IMAGE_DISPONIBLE or PYMUPDF_DISPONIBLE,
            "guardado_automatico": guardado_automatico,
            "carpeta_guardado": str(self.carpeta_fecha)
        }
    
    async def procesar_multiples_archivos(self, archivos: list) -> Dict[str, str]:
        """
        Procesa múltiples archivos y retorna diccionario con textos extraídos.
        GUARDA AUTOMÁTICAMENTE cada extracción individual.
        
        Args:
            archivos: Lista de archivos UploadFile
            
        Returns:
            Dict[str, str]: Diccionario {nombre_archivo: texto_extraido}
        """
        textos_extraidos = {}
        
        logger.info(f" Procesando {len(archivos)} archivos con guardado automático")
        logger.info(f" Carpeta de guardado: {self.carpeta_fecha}")
        
        for archivo in archivos:
            try:
                # Validar archivo
                validacion = self.validar_archivo(archivo)
                if not validacion["valido"]:
                    logger.error(f" Archivo inválido {archivo.filename}: {validacion['error']}")
                    textos_extraidos[archivo.filename] = f"ERROR: {validacion['error']}"
                    continue
                
                # Procesar archivo (esto automáticamente guarda el texto)
                texto = await self.procesar_archivo(archivo)
                textos_extraidos[archivo.filename] = texto
                
                logger.info(f" Archivo procesado y guardado: {archivo.filename}")
                
            except Exception as e:
                logger.error(f" Error procesando archivo {archivo.filename}: {e}")
                textos_extraidos[archivo.filename] = f"ERROR PROCESANDO: {str(e)}"
                
                # Guardar también los errores de procesamiento
                self._guardar_texto_extraido(
                    archivo.filename, f"ERROR PROCESANDO: {str(e)}", 
                    "PROCESAMIENTO_ERROR", {"error": str(e)}
                )
        
        logger.info(f" Procesamiento completado: {len(textos_extraidos)} archivos")
        logger.info(f" Todos los textos extraídos guardados en: {self.carpeta_fecha}")
        
        return textos_extraidos
    
    def obtener_estadisticas_guardado(self) -> Dict[str, Any]:
        """
        Obtiene estadísticas de los archivos guardados hoy.
        
        Returns:
            Dict con estadísticas de guardado
        """
        try:
            archivos_guardados = list(self.carpeta_fecha.glob("*.txt"))
            
            estadisticas = {
                "fecha": datetime.now().strftime("%Y-%m-%d"),
                "carpeta": str(self.carpeta_fecha),
                "total_archivos_guardados": len(archivos_guardados),
                "tipos_extraccion": {},
                "tamaño_total_mb": 0,
                "dependencias": {
                    "google_vision": self.vision_client is not None,
                    "pdfplumber": PDFPLUMBER_DISPONIBLE,
                    "pdf2image": PDF2IMAGE_DISPONIBLE,
                    "pymupdf": PYMUPDF_DISPONIBLE,
                    "extract_msg": EXTRACT_MSG_DISPONIBLE
                }
            }
            
            # Analizar archivos guardados
            for archivo in archivos_guardados:
                # Extraer tipo de extracción del nombre
                partes = archivo.name.split("_")
                if len(partes) >= 2:
                    tipo = partes[1]
                    estadisticas["tipos_extraccion"][tipo] = estadisticas["tipos_extraccion"].get(tipo, 0) + 1
                
                # Sumar tamaño
                estadisticas["tamaño_total_mb"] += archivo.stat().st_size / (1024 * 1024)
            
            estadisticas["tamaño_total_mb"] = round(estadisticas["tamaño_total_mb"], 2)
            
            return estadisticas
            
        except Exception as e:
            logger.error(f"Error obteniendo estadísticas: {e}")
            return {"error": str(e)}


# ===============================
# FUNCIONES DE PREPROCESAMIENTO EXCEL
# ===============================

def preprocesar_excel_limpio(contenido: bytes, nombre_archivo: str = "archivo.xlsx") -> str:
    """
    Preprocesa archivo Excel eliminando filas y columnas vacías.
    Mantiene formato tabular limpio con toda la información intacta.

    FUNCIONALIDAD:
    Elimina filas completamente vacías
    Elimina columnas completamente vacías
    Mantiene formato tabular pero limpio
    Conserva toda la información relevante
    Óptimo y simple
    Guarda automáticamente el archivo preprocesado

    Args:
        contenido: Contenido binario del archivo Excel
        nombre_archivo: Nombre del archivo (para logging)

    Returns:
        str: Texto extraído y limpio del Excel
    """
    try:
        logger.info(f" Preprocesando Excel: {nombre_archivo}")

        # 1. LEER EXCEL CON TODAS LAS HOJAS
        df_dict = pd.read_excel(io.BytesIO(contenido), sheet_name=None)

        texto_completo = ""
        total_hojas = 0
        filas_eliminadas_total = 0
        columnas_eliminadas_total = 0

        # 2. PROCESAR CADA HOJA CON LIMPIEZA
        if isinstance(df_dict, dict):
            total_hojas = len(df_dict)

            for nombre_hoja, dataframe in df_dict.items():
                # Estadísticas originales
                filas_orig = len(dataframe)
                cols_orig = len(dataframe.columns)

                #  LIMPIEZA SIMPLE: Eliminar filas y columnas completamente vacías
                df_limpio = dataframe.dropna(how='all')  # Filas vacías
                df_limpio = df_limpio.dropna(axis=1, how='all')  # Columnas vacías

                # Estadísticas después de limpieza
                filas_final = len(df_limpio)
                cols_final = len(df_limpio.columns)

                filas_eliminadas = filas_orig - filas_final
                columnas_eliminadas = cols_orig - cols_final
                filas_eliminadas_total += filas_eliminadas
                columnas_eliminadas_total += columnas_eliminadas

                # Agregar hoja al texto
                texto_completo += f"\n--- HOJA: {nombre_hoja} ---\n"

                if not df_limpio.empty:
                    # Convertir a texto manteniendo formato tabular limpio
                    texto_hoja = df_limpio.to_string(index=False, na_rep='', max_cols=None, max_rows=None)
                    texto_completo += texto_hoja
                else:
                    texto_completo += "[HOJA VACÍA DESPUÉS DE LIMPIEZA]"

                texto_completo += "\n"

        else:
            # UNA SOLA HOJA
            total_hojas = 1
            dataframe = df_dict

            # Estadísticas originales
            filas_orig = len(dataframe)
            cols_orig = len(dataframe.columns)

            #  LIMPIEZA SIMPLE: Eliminar filas y columnas vacías
            df_limpio = dataframe.dropna(how='all')  # Filas vacías
            df_limpio = df_limpio.dropna(axis=1, how='all')  # Columnas vacías

            # Estadísticas finales
            filas_final = len(df_limpio)
            cols_final = len(df_limpio.columns)

            filas_eliminadas_total = filas_orig - filas_final
            columnas_eliminadas_total = cols_orig - cols_final

            if not df_limpio.empty:
                texto_completo = df_limpio.to_string(index=False, na_rep='', max_cols=None, max_rows=None)
            else:
                texto_completo = "[ARCHIVO VACÍO DESPUÉS DE LIMPIEZA]"

        texto_final = texto_completo.strip()

        # 3. GUARDADO AUTOMÁTICO DEL ARCHIVO PREPROCESADO
        _guardar_archivo_preprocesado(nombre_archivo, texto_final, filas_eliminadas_total, columnas_eliminadas_total, total_hojas)

        # 4. LOGGING OPTIMIZADO
        logger.info(f" Preprocesamiento completado: {len(texto_final)} caracteres")
        logger.info(f" Hojas: {total_hojas} | Filas eliminadas: {filas_eliminadas_total} | Columnas eliminadas: {columnas_eliminadas_total}")
        logger.info(f" Archivo preprocesado guardado automáticamente")

        return texto_final

    except Exception as e:
        error_msg = f"Error en preprocesamiento Excel: {str(e)}"
        logger.error(f" {error_msg}")
        return error_msg

def _guardar_archivo_preprocesado(nombre_archivo: str, texto_preprocesado: str,
                                 filas_eliminadas: int, columnas_eliminadas: int, total_hojas: int):
    """
    Guarda el archivo Excel preprocesado según nomenclatura {archivo_original}_preprocesado.txt

    FUNCIONALIDAD:
     Guarda en carpeta extracciones/
     Nomenclatura: {archivo_original}_preprocesado.txt
     Logs básicos para confirmar guardado exitoso
     Manejo de errores sin afectar flujo principal

    Args:
        nombre_archivo: Nombre del archivo original
        texto_preprocesado: Texto limpio extraído
        filas_eliminadas: Número de filas eliminadas
        columnas_eliminadas: Número de columnas eliminadas
        total_hojas: Número total de hojas procesadas
    """
    try:
        # 1. CREAR CARPETA EXTRACCIONES SIMPLE
        carpeta_extracciones = Path("extracciones")
        carpeta_extracciones.mkdir(exist_ok=True)

        # 2. CREAR NOMBRE SEGÚN NOMENCLATURA: {archivo_original}_preprocesado.txt
        # Limpiar nombre de archivo original (quitar caracteres especiales)
        nombre_base = "".join(c for c in nombre_archivo if c.isalnum() or c in "._-")

        # Quitar extensión original (.xlsx, .xls)
        if '.' in nombre_base:
            nombre_sin_extension = nombre_base.rsplit('.', 1)[0]
        else:
            nombre_sin_extension = nombre_base

        # Crear nombre final: {archivo_original}_preprocesado.txt
        nombre_final = f"{nombre_sin_extension}_preprocesado.txt"
        ruta_archivo = carpeta_extracciones / nombre_final

        # 3. CONTENIDO SIMPLE PERO COMPLETO
        contenido_final = f"""ARCHIVO EXCEL PREPROCESADO
=============================

Archivo original: {nombre_archivo}
Fecha procesamiento: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
Hojas procesadas: {total_hojas}
Filas vacías eliminadas: {filas_eliminadas}
Columnas vacías eliminadas: {columnas_eliminadas}
Caracteres finales: {len(texto_preprocesado)}

=============================
TEXTO ENVIADO A GEMINI:
=============================

{texto_preprocesado}
"""

        # 4. GUARDAR ARCHIVO
        with open(ruta_archivo, 'w', encoding='utf-8') as f:
            f.write(contenido_final)

        # 5. LOG BÁSICO DE CONFIRMACIÓN
        logger.info(f" Archivo preprocesado guardado: extracciones/{nombre_final}")
        logger.info(f" Estadísticas: {filas_eliminadas} filas y {columnas_eliminadas} columnas eliminadas")

    except Exception as e:
        logger.error(f" Error guardando archivo preprocesado: {e}")
        # No fallar el preprocesamiento por un error de guardado
